"""
Gera o Artefato 1 - Customer Segmentation Analysis Report (Word) com
visualizacoes, descricao dos clusters, hipoteses de drivers de churn,
priorizacao de intervencao e secao de limitacoes.

Saida: Artefato_1_Segmentation_Report.docx
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
REPORT_JSON = BASE_DIR / "segmentation_report.json"
CHARTS_DIR = BASE_DIR / "charts"
OUT_PATH = BASE_DIR / "Artefato_1_Segmentation_Report.docx"

PRIMARY = RGBColor(0x26, 0x46, 0x53)
ACCENT = RGBColor(0x2A, 0x9D, 0x8F)
DANGER = RGBColor(0xE7, 0x6F, 0x51)


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1, color: RGBColor = PRIMARY) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_image_with_caption(doc: Document, path: Path, caption: str, width_cm: float = 14) -> None:
    doc.add_picture(str(path), width=Cm(width_cm))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build_cluster_table(doc: Document, profiles: list[dict], names: dict[str, dict]) -> None:
    headers = ["Cluster", "Nome", "Tamanho", "Share", "Churn %", "Lift vs base"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "264653")

    for p in profiles:
        row = table.add_row().cells
        row[0].text = f"C{p['cluster']}"
        row[1].text = names[str(p["cluster"])]["nome"]
        row[2].text = f"{p['size']:,}".replace(",", ".")
        row[3].text = f"{p['share']:.1%}"
        row[4].text = f"{p['churn_rate']:.1%}"
        row[5].text = f"{p['lift_vs_base']:.2f}x"

        churn_color = "f4a261" if p["churn_rate"] >= 0.45 else (
            "e9c46a" if p["churn_rate"] >= 0.20 else "d9ed92"
        )
        set_cell_shading(row[4], churn_color)
        for c in row:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_priority_table(doc: Document, priority: list[dict]) -> None:
    headers = ["#", "Cluster", "Share da base", "Churn", "Score", "Acao recomendada"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "264653")

    for rank, p in enumerate(priority, start=1):
        row = table.add_row().cells
        row[0].text = str(rank)
        row[1].text = f"C{p['cluster']} - {p['nome']}"
        row[2].text = f"{p['share']:.1%}"
        row[3].text = f"{p['churn_rate']:.1%}"
        row[4].text = f"{p['score_prioridade']:.3f}"
        row[5].text = p["acao_recomendada"]
        if rank == 1:
            set_cell_shading(row[0], "e76f51")
        elif rank == 2:
            set_cell_shading(row[0], "f4a261")
        elif rank == 3:
            set_cell_shading(row[0], "e9c46a")


def build() -> None:
    with open(REPORT_JSON, "r", encoding="utf-8") as f:
        data = json.load(f)

    profiles = data["profiles"]
    names = data["names"]
    priority = data["priority"]
    extras = data["extras"]

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # CAPA
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ARTEFATO 1")
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    run.bold = True

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("Customer Segmentation Analysis Report")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Case Vitaliza | Modulo 2 - Inteli MBA | Turma 2026.1")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # SUMARIO EXECUTIVO
    add_heading(doc, "1. Sumario executivo", level=1)
    add_paragraph(
        doc,
        "Esta analise segmenta a base de 4.000 clientes Vitaliza em 4 clusters via "
        "KMeans, usando as 13 variaveis comportamentais e contratuais do dataset. O "
        "objetivo e identificar QUEM cancela, POR QUE cancela e ONDE a intervencao "
        "tem maior retorno marginal antes do Board de 14 de novembro.",
    )
    add_paragraph(doc, "Achados centrais:", bold=True)
    add_bullets(
        doc,
        [
            f"Taxa de churn global na base: {extras['churn_global']:.1%}. Tres "
            "drivers concentram quase toda a variancia: contrato curto, lifetime baixo "
            "e queda na frequencia do mes corrente.",
            "O cluster de maior risco (Early Dropout) representa 36,8% da base mas "
            "responde por mais de 78% dos cancelamentos previstos - e o maior alvo de ROI.",
            "Clientes em contrato anual cancelam 18x menos que clientes em contrato "
            "mensal (2,4% vs 42,3%). A duracao do contrato e a alavanca contratual mais "
            "obvia para reduzir churn.",
            "Clientes que entraram por indique-um-amigo ou frequentam aulas em grupo "
            "tem churn quase 50% menor. Sao sinais sociais que o modelo captura.",
        ],
    )

    # METODOLOGIA
    add_heading(doc, "2. Metodologia", level=1)
    add_bullets(
        doc,
        [
            "Dataset: gym_churn_us.csv (4.000 linhas, 14 colunas). Target = Churn.",
            "Pre-processamento: StandardScaler nas 13 features numericas (sem encoding "
            "adicional - todas ja sao binarias ou numericas).",
            "Clusterizacao: KMeans com k=4, n_init=20, random_state=42. Clusters "
            "reordenados pela taxa de churn (C0 = menor risco, C3 = maior risco).",
            "Profiling: media de cada feature por cluster + taxa de churn observada + "
            "lift versus a base.",
            "Priorizacao: score = share x churn_rate (intuicao de 'volume absoluto de "
            "cancelamentos atribuiveis ao cluster').",
        ],
    )

    # EDA - 4 graficos
    add_heading(doc, "3. Analise exploratoria", level=1)
    add_paragraph(
        doc,
        "Antes de segmentar, isolamos as variaveis com correlacao mais forte com churn. "
        "Os graficos abaixo mostram que o sinal nao esta uniformemente distribuido: "
        "ele se concentra em contratos curtos, lifetime baixo e queda de frequencia.",
    )
    add_image_with_caption(doc, CHARTS_DIR / "01_churn_rate.png",
                           "Figura 1. Distribuicao geral - 26,5% da base esta marcada como churn.")
    add_image_with_caption(doc, CHARTS_DIR / "02_churn_by_contract.png",
                           "Figura 2. Churn por duracao do contrato. Contrato mensal tem 42% de churn; anual, 2%.")
    add_image_with_caption(doc, CHARTS_DIR / "03_churn_by_lifetime.png",
                           "Figura 3. Churn por lifetime. Clientes com menos de 1 mes concentram o maior risco.")
    add_image_with_caption(doc, CHARTS_DIR / "04_churn_by_frequency.png",
                           "Figura 4. Churn por frequencia media no mes corrente - virtualmente zero acima de 3,5 aulas/semana.")
    add_image_with_caption(doc, CHARTS_DIR / "05_correlation.png",
                           "Figura 5. Correlacao linear de cada feature com Churn. Lifetime, Contract_period "
                           "e frequencia atual sao os 3 sinais mais fortes (negativos).")

    # SEGMENTACAO
    add_heading(doc, "4. Resultado da segmentacao", level=1)
    add_paragraph(
        doc,
        "Quatro clusters emergem do KMeans com perfis claramente distintos em "
        "engajamento, contrato e lifetime. A tabela abaixo resume tamanho e risco "
        "de cada um.",
    )
    build_cluster_table(doc, profiles, names)
    doc.add_paragraph()
    add_image_with_caption(doc, CHARTS_DIR / "06_cluster_sizes.png",
                           "Figura 6. Distribuicao da base pelos 4 segmentos.")
    add_image_with_caption(doc, CHARTS_DIR / "07_cluster_churn_rates.png",
                           "Figura 7. Taxa de churn observada em cada segmento (linha = media global 26,5%).")
    add_image_with_caption(doc, CHARTS_DIR / "08_cluster_radar.png",
                           "Figura 8. Perfil comparativo (normalizado) dos clusters em 6 dimensoes-chave.", width_cm=12)

    # DESCRICAO DETALHADA POR CLUSTER
    add_heading(doc, "5. Descricao dos clusters e hipoteses de drivers", level=1)
    for p in sorted(profiles, key=lambda x: -x["churn_rate"]):
        cid = p["cluster"]
        meta = names[str(cid)]
        m = p["means"]
        add_heading(doc, f"5.{4 - cid} Cluster C{cid} - {meta['nome']}", level=2)
        add_paragraph(doc, meta["descricao"], italic=True)
        add_bullets(
            doc,
            [
                f"Tamanho: {p['size']:,} clientes ({p['share']:.1%} da base)".replace(",", "."),
                f"Churn observado: {p['churn_rate']:.1%} (lift {p['lift_vs_base']:.2f}x sobre a media).",
                f"Perfil tipico: contrato medio de {m['Contract_period']:.1f} meses, lifetime "
                f"de {m['Lifetime']:.1f} meses, frequencia atual de "
                f"{m['Avg_class_frequency_current_month']:.2f} aulas/semana, idade "
                f"{m['Age']:.0f} anos.",
                f"Spend medio em servicos extras: USD {m['Avg_additional_charges_total']:.0f}.",
                f"Aderencia a aulas em grupo: {m['Group_visits']*100:.0f}%; "
                f"Promo amigo: {m['Promo_friends']*100:.0f}%; "
                f"Parceria corporativa: {m['Partner']*100:.0f}%.",
            ],
        )
        add_paragraph(doc, "Hipoteses sobre o driver de churn deste cluster:", bold=True)
        if cid == 3:  # Early dropout
            add_bullets(doc, [
                "H1: Onboarding falho - clientes que chegam ao mes 2 sem frequencia minima "
                "de 1,5 aulas/semana tem probabilidade >50% de cancelar. Drive: ausencia de "
                "ritual de adocao nas primeiras 4 semanas.",
                "H2: Mismatch de plano - contratos mensais sao porta de entrada e tambem de "
                "saida; o cliente nao se compromete e sai na primeira friccao.",
                "H3: Isolamento social - <30% participam de aulas em grupo; quem participa "
                "tem churn metade menor. O vinculo social esta ausente.",
            ])
        elif cid == 2:  # Risco medio
            add_bullets(doc, [
                "H1: Clientes 'no limbo' - frequentam, mas a frequencia do mes corrente caiu "
                "vs a media historica. Sinal claro de desengajamento gradual.",
                "H2: Contrato semestral ainda em vigor cria atraso na decisao - o churn vai "
                "se materializar quando o contrato terminar (Month_to_end_contract baixo).",
                "H3: Hipotese 'sleeping dog' parcial - intervir errado pode acelerar a saida.",
            ])
        elif cid == 1:  # Engajados mensais
            add_bullets(doc, [
                "H1: Baixo churn (~9%) mas exposto a saida por motivos exogenos (preco, "
                "concorrencia, mudanca de rotina). Sao alvos naturais de upsell para "
                "contrato semestral/anual.",
                "H2: Frequencia alta e estavel ao longo dos meses - o engajamento e real, "
                "nao inflacionado por evento unico.",
            ])
        else:  # Leais contrato longo
            add_bullets(doc, [
                "H1: Cluster 'imune' no horizonte de 1 mes - churn de 3% e dentro do ruido. "
                "Atuar aqui tem retorno marginal baixo.",
                "H2: Risco real e perda na renovacao do contrato anual. Monitorar "
                "Month_to_end_contract para acionar oferta de renovacao 60 dias antes.",
            ])

    # PRIORIZACAO
    add_heading(doc, "6. Priorizacao de intervencao", level=1)
    add_paragraph(
        doc,
        "Combinando tamanho do cluster com taxa de churn, obtemos um score de "
        "prioridade que estima quantos cancelamentos cada segmento contribui para o "
        "total. A acao recomendada segue a logica de Caminho A (reativo) + Caminho B "
        "(proativo) discutida no case.",
    )
    build_priority_table(doc, priority)
    doc.add_paragraph()
    add_paragraph(doc, "Leitura executiva:", bold=True)
    add_bullets(doc, [
        "C3 (Early Dropout) e onde 70-80% do esforco deve estar concentrado nas "
        "primeiras 4 semanas - maior volume e maior taxa.",
        "C2 (Risco Medio) entra na cesta de monitoramento. Modelo Caminho B com "
        "intervencoes leves (email com programa novo) - sem desperdicar oferta agressiva.",
        "C1 (Engajados Mensais) deve receber upsell para contrato semestral - reduz "
        "churn estrutural sem ofertar desconto.",
        "C0 (Leais) NAO deve receber intervencao ativa - risco de efeito sleeping dog.",
    ])

    # LIMITACOES
    add_heading(doc, "7. Limitacoes da analise", level=1)
    add_paragraph(
        doc,
        "Esta secao e exigida pela orientacao do artefato e e onde o stakeholder "
        "tecnico (CTO Diego) precisa ser ouvido. Sao limites materiais que "
        "comprometem a generalizacao.",
        italic=True,
    )
    add_bullets(doc, [
        "Ausencia de dados granulares do Mixpanel: o dataset agregado nao permite "
        "construir features de serie temporal (declinio de frequencia semana a semana, "
        "tempo desde ultima sessao, sequencia de eventos). A hipotese 'declinio de "
        "engajamento prevê churn' nao pode ser testada com o que temos.",
        "Impossibilidade de analisar serie temporal de engajamento: as variaveis "
        "Avg_class_frequency_total e Avg_class_frequency_current_month sao snapshots, "
        "nao trajetorias. Perdemos sinais de momentum.",
        "Churn invisivel nao representado: cancelamentos via falha de pagamento, "
        "chargeback ou abandono silencioso (estimado em ~33% no case original) nao "
        "estao rotulados no dataset. Qualquer modelo treinado aqui tem recall maximo "
        "teorico de ~67% do fenomeno total.",
        "Possivel target leakage: Avg_class_frequency_current_month e "
        "Month_to_end_contract sao co-temporais ao evento. Para producao, e necessario "
        "reconstruir features em snapshot point-in-time (T-30 dias).",
        "Sem holdout temporal: o split aleatorio sub-estima a deterioracao do modelo "
        "no tempo (drift). Recomendado teste com split temporal (treino jan-jun, teste "
        "jul-set) antes do deploy.",
        "Variaveis ausentes que importariam: motivo de cancelamento, NPS, ticket de "
        "suporte aberto, fonte de aquisicao (organico vs paid), historico de pagamento.",
    ])

    # PROXIMOS PASSOS
    add_heading(doc, "8. Proximos passos para o time de dados", level=1)
    add_bullets(doc, [
        "Semanas 1-3: reconciliar PostgreSQL x Mixpanel x GA4 por user_id. Sem isso, "
        "C2 e C3 permanecem indistinguiveis em producao.",
        "Semanas 3-5: reconstruir snapshot point-in-time T-30 e retreinar o modelo "
        "preditivo (ja temos GradientBoosting AUC=0,98 no benchmark atual, mas com risco "
        "de leakage).",
        "Semanas 5-7: integrar a segmentacao a interface (este artefato) para que CS e "
        "Marketing recebam o cluster do cliente em tempo real.",
        "Semanas 8-10: piloto A/B - intervir em 50% do C3 com Caminho B (proativo) e "
        "comparar churn vs grupo controle. Validar que a intervencao CAUSA retencao.",
    ])

    doc.save(OUT_PATH)
    print(f"OK -> {OUT_PATH}")


if __name__ == "__main__":
    build()
